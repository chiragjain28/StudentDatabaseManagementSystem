# import smtplib
# server = smtplib.SMTP('smtp.gmail.com', 587)
# server.ehlo()
# server.starttls()
# server.login('cjain3631@gmail.com', 'Chiragjain3631')
# SUBJECT="Congratulations"
# TEXT="Your Email ans password is "
# message = 'Subject: {}\n\n{}'.format(SUBJECT, TEXT)
# server.sendmail('cjain3631@gmail.com','cjain3631@gmail.com' , message)
# server.close()

# self.sendEmail(self.var_email.get(),self.var_name.get(),self.var_pass.get(),self.var_course.get())
# f'''
# Dear {name[0]}!

# Congratulations!

# You’ve successfully registered for {course}. Now you can access our 'Imposter's portal' 
# and update your information

# Login id : {to}
# Password : {pas}

# IMPORTANT NOTICE : 

# "Don't worry about the course fee and expenses , Due to this pandemic, we are halving the fees, because you are not access any college things, you can pay the fees anytime.
# We understand your compulsion " 

# If you have any questions or concerns, feel free to contact me at cjain3631@gmail.com or 7828872160.

# Best,
# Imposters


# "Dear "+name+"!"+"\nCongratulations !"+ "\nYou’ve successfully registered for cousre "+course+"\nNow you can access our 'Imposter's portal' and update your information"+"\nYour login details :"+"\nLogin id : "+to+"\nPassword : "+pas+" IMPORTANT NOTICE :\n

# Don't worry about the course fee and expenses , Due to this pandemic, We are halving the fees,\nbecause you are not access any college things,and you can pay the fees anytime.\nWe understand your compulsion\n\n
# If you have any questions or concerns, feel free to contact me at cjain3631@gmail.com or 7828872160.\n\n\nBest,Imposters"    

# '''
import docx
from docx.enum.text import WD_COLOR_INDEX
  
# Create an instance of a word document
doc = docx.Document()
highlight_para = doc.add_paragraph(
       ).add_run(
           '''GeeksforGeeks is a Computer Science portal for geeks. It contains well written, well thought and well-explained computer science and programming articles, quizzes etc.'''
                 ).font.highlight_color = WD_COLOR_INDEX.YELLOW